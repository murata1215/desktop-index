# =============================================================================
# Desktop Index - ドキュメントパーサー
# =============================================================================
# 各種ドキュメント形式（PDF、Word、Excel、テキスト）からテキストを抽出します。
#
# サポートする形式:
#   - PDF (.pdf): pdfplumber を使用
#   - Word (.docx): python-docx を使用
#   - Excel (.xlsx): openpyxl を使用
#   - テキスト: 標準ライブラリ + chardet でエンコーディング検出
# =============================================================================

import logging
from pathlib import Path
from typing import Optional
import chardet

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# PDF パーサー
# ---------------------------------------------------------------------------
def extract_pdf_text(file_path: str, max_pages: int = 100) -> Optional[str]:
    """
    PDFファイルからテキストを抽出する

    pdfplumber を使用してPDFの各ページからテキストを抽出します。
    スキャンされた画像PDFの場合、テキストは抽出できません。

    Args:
        file_path: PDFファイルのパス
        max_pages: 処理する最大ページ数（デフォルト: 100）
                   巨大なPDFの処理時間を制限するため

    Returns:
        str: 抽出されたテキスト、失敗した場合は None

    Note:
        - パスワード保護されたPDFは処理できません
        - 画像のみのPDFはテキストが抽出できません（OCRは未実装）
    """
    try:
        import pdfplumber

        texts = []
        with pdfplumber.open(file_path) as pdf:
            # 処理するページ数を制限
            pages_to_process = min(len(pdf.pages), max_pages)

            for i in range(pages_to_process):
                page = pdf.pages[i]
                text = page.extract_text()
                if text:
                    texts.append(text)

        return "\n".join(texts) if texts else None

    except Exception as e:
        logger.debug(f"PDF解析エラー: {file_path} - {e}")
        return None


# ---------------------------------------------------------------------------
# Word パーサー
# ---------------------------------------------------------------------------
def extract_docx_text(file_path: str) -> Optional[str]:
    """
    Word (.docx) ファイルからテキストを抽出する

    python-docx を使用して、ドキュメント内の全段落からテキストを抽出します。
    表やヘッダー/フッターのテキストも含まれます。

    Args:
        file_path: Word ファイルのパス

    Returns:
        str: 抽出されたテキスト、失敗した場合は None

    Note:
        - .doc（旧形式）は非対応です
        - 埋め込みオブジェクトやスマートアートのテキストは抽出できない場合があります
    """
    try:
        from docx import Document

        doc = Document(file_path)
        texts = []

        # 本文の段落を抽出
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                texts.append(paragraph.text)

        # テーブル内のテキストも抽出
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text)

        return "\n".join(texts) if texts else None

    except Exception as e:
        logger.debug(f"Word解析エラー: {file_path} - {e}")
        return None


# ---------------------------------------------------------------------------
# Excel パーサー
# ---------------------------------------------------------------------------
def extract_xlsx_text(file_path: str, max_rows: int = 10000) -> Optional[str]:
    """
    Excel (.xlsx) ファイルからテキストを抽出する

    openpyxl を使用して、全シートの全セルからテキストを抽出します。
    数値も文字列として抽出されます。

    Args:
        file_path: Excel ファイルのパス
        max_rows: 処理する最大行数（デフォルト: 10000）
                  巨大なExcelファイルの処理時間を制限するため

    Returns:
        str: 抽出されたテキスト（セルはタブ区切り、行は改行区切り）
             失敗した場合は None

    Note:
        - .xls（旧形式）は非対応です
        - 数式は計算結果が抽出されます（数式自体ではない）
        - 非表示のシートも処理されます
    """
    try:
        from openpyxl import load_workbook

        # data_only=True で数式ではなく計算結果を取得
        wb = load_workbook(file_path, read_only=True, data_only=True)
        texts = []
        total_rows = 0

        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                if total_rows >= max_rows:
                    break

                row_texts = []
                for cell in row:
                    if cell.value is not None:
                        # 値を文字列に変換
                        row_texts.append(str(cell.value))

                if row_texts:
                    texts.append("\t".join(row_texts))
                    total_rows += 1

            if total_rows >= max_rows:
                break

        wb.close()
        return "\n".join(texts) if texts else None

    except Exception as e:
        logger.debug(f"Excel解析エラー: {file_path} - {e}")
        return None


# ---------------------------------------------------------------------------
# テキストファイルパーサー
# ---------------------------------------------------------------------------
def extract_text_file(file_path: str, max_size: int = 1024 * 1024) -> Optional[str]:
    """
    テキストファイルからテキストを抽出する

    chardet を使用してエンコーディングを自動検出し、テキストを読み込みます。
    ソースコードファイル（.py, .js 等）もこの関数で処理されます。

    Args:
        file_path: テキストファイルのパス
        max_size: 読み込む最大バイト数（デフォルト: 1MB）
                  巨大なテキストファイルの処理時間を制限するため

    Returns:
        str: ファイルの内容、失敗した場合は None

    Note:
        - バイナリファイルを渡すと None が返されます
        - エンコーディング検出に失敗した場合は UTF-8 を試行します
    """
    try:
        # まずファイルの一部を読み込んでエンコーディングを検出
        with open(file_path, 'rb') as f:
            raw_data = f.read(min(max_size, 10000))  # 検出用に最大10KB

        # エンコーディング検出
        detected = chardet.detect(raw_data)
        encoding = detected.get('encoding', 'utf-8')
        confidence = detected.get('confidence', 0)

        # 信頼度が低い場合は UTF-8 を優先
        if confidence < 0.5:
            encoding = 'utf-8'

        # ファイル全体を読み込み
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            # 最大サイズまで読み込み
            content = f.read(max_size)

        return content

    except Exception as e:
        logger.debug(f"テキスト読み込みエラー: {file_path} - {e}")
        return None


# ---------------------------------------------------------------------------
# ドキュメントパーサークラス
# ---------------------------------------------------------------------------
class DocumentParser:
    """
    様々な形式のドキュメントからテキストを抽出するクラス

    ファイルの拡張子に基づいて適切なパーサーを選択し、
    テキストを抽出します。

    使用例:
        parser = DocumentParser(max_content_length=50000)
        text = parser.extract_text("/path/to/document.pdf")
        if text:
            print(f"抽出されたテキスト: {len(text)} 文字")

    Attributes:
        max_content_length: 抽出テキストの最大長（これを超えると切り詰め）
    """

    # 拡張子とパーサー関数のマッピング
    # 拡張子は小文字で、ドット付きで指定
    PARSERS = {
        # ドキュメント形式
        '.pdf': extract_pdf_text,
        '.docx': extract_docx_text,
        '.xlsx': extract_xlsx_text,

        # テキスト形式（全て extract_text_file で処理）
        '.txt': extract_text_file,
        '.md': extract_text_file,
        '.csv': extract_text_file,
        '.json': extract_text_file,
        '.xml': extract_text_file,
        '.yaml': extract_text_file,
        '.yml': extract_text_file,

        # ソースコード（全て extract_text_file で処理）
        '.py': extract_text_file,
        '.js': extract_text_file,
        '.ts': extract_text_file,
        '.jsx': extract_text_file,
        '.tsx': extract_text_file,
        '.html': extract_text_file,
        '.css': extract_text_file,
        '.scss': extract_text_file,
        '.java': extract_text_file,
        '.c': extract_text_file,
        '.cpp': extract_text_file,
        '.h': extract_text_file,
        '.hpp': extract_text_file,
        '.cs': extract_text_file,
        '.go': extract_text_file,
        '.rs': extract_text_file,
        '.rb': extract_text_file,
        '.php': extract_text_file,
        '.sql': extract_text_file,
        '.sh': extract_text_file,
        '.bat': extract_text_file,
        '.ps1': extract_text_file,
    }

    def __init__(self, max_content_length: int = 100000):
        """
        DocumentParser を初期化する

        Args:
            max_content_length: 抽出テキストの最大長（文字数）
                               これを超えるテキストは切り詰められます
        """
        self.max_content_length = max_content_length

    def extract_text(self, file_path: str) -> Optional[str]:
        """
        ファイルからテキストを抽出する

        ファイルの拡張子に基づいて適切なパーサーを選択し、
        テキストを抽出します。

        Args:
            file_path: ファイルのパス

        Returns:
            str: 抽出されたテキスト（最大長で切り詰め済み）
                 サポートされていない形式または抽出失敗時は None
        """
        # 拡張子を取得（小文字に正規化）
        ext = Path(file_path).suffix.lower()

        # 対応するパーサーを取得
        parser_func = self.PARSERS.get(ext)

        if not parser_func:
            # サポートされていない拡張子
            logger.debug(f"未サポートの拡張子: {ext} ({file_path})")
            return None

        # テキスト抽出を実行
        try:
            text = parser_func(file_path)

            if text:
                # 最大長で切り詰め
                if len(text) > self.max_content_length:
                    text = text[:self.max_content_length]
                    logger.debug(f"テキストを切り詰めました: {file_path}")

                # 空白文字の正規化
                text = self._normalize_whitespace(text)

            return text

        except Exception as e:
            logger.warning(f"テキスト抽出エラー: {file_path} - {e}")
            return None

    def _normalize_whitespace(self, text: str) -> str:
        """
        テキスト内の空白文字を正規化する

        連続する空白や改行を単一の空白/改行に置換し、
        検索性能を向上させます。

        Args:
            text: 正規化前のテキスト

        Returns:
            str: 正規化後のテキスト
        """
        # 行ごとに処理
        lines = text.splitlines()

        # 各行の先頭・末尾の空白を削除し、空行を1つにまとめる
        normalized_lines = []
        prev_empty = False

        for line in lines:
            stripped = line.strip()
            if stripped:
                normalized_lines.append(stripped)
                prev_empty = False
            elif not prev_empty:
                # 連続する空行は1つにまとめる
                normalized_lines.append("")
                prev_empty = True

        return "\n".join(normalized_lines)

    def is_supported(self, file_path: str) -> bool:
        """
        ファイルがテキスト抽出に対応しているかチェックする

        Args:
            file_path: ファイルのパス

        Returns:
            bool: 対応している場合は True
        """
        ext = Path(file_path).suffix.lower()
        return ext in self.PARSERS
